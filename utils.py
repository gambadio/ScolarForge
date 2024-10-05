# utils.py

import json
import os
import tkinter as tk
from tkinter import filedialog, messagebox
import requests
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import PyPDF2
import re
from bs4 import BeautifulSoup

class FileHandler:
    def get_file_paths(self, title):
        file_types = [("PDF files", "*.pdf"), ("Text files", "*.txt"), ("All files", "*.*")]
        return filedialog.askopenfilenames(title=title, filetypes=file_types)

    def upload_script(self, parent, file_path):
        file_name = os.path.basename(file_path)
        text = self.extract_text_from_file(file_path)
        parent.scripts.append((file_name, text))
        self.save_script_texts(parent)

    def upload_instruction(self, parent, file_path):
        file_name = os.path.basename(file_path)
        text = self.extract_text_from_file(file_path)
        parent.instructions.append((file_name, text))
        self.save_instruction_texts(parent)

    def extract_text_from_file(self, file_path):
        file_extension = os.path.splitext(file_path)[1].lower()
        if file_extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        else:
            return self.extract_text_from_txt(file_path)

    def extract_text_from_pdf(self, file_path):
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                return "".join(page.extract_text() for page in reader.pages)
        except Exception as e:
            messagebox.showerror("Error", f"Error reading PDF file: {e}")
            return ""

    def extract_text_from_txt(self, file_path):
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    return f.read()
            except Exception as e:
                messagebox.showerror("Error", f"Error reading text file: {e}")
                return ""

    def scrape_webpage(self, url):
        try:
            response = requests.get(url)
            response.raise_for_status()
            soup = BeautifulSoup(response.content, 'html.parser')
            text = soup.get_text(separator='\n')
            return text.strip()
        except requests.RequestException as e:
            messagebox.showerror("Error", f"Error fetching web page: {e}")
            return None

    def save_script_texts(self, parent):
        self.save_texts(parent.scripts, 'script_texts.json')

    def save_instruction_texts(self, parent):
        self.save_texts(parent.instructions, 'instruction_texts.json')

    def save_internet_sources(self, parent):
        try:
            with open('internet_sources.json', 'w') as f:
                json.dump(parent.internet_sources, f, indent=2)
        except Exception as e:
            messagebox.showerror("Error", f"Error saving internet sources: {e}")

    def save_internet_search_results(self, parent):
        try:
            with open('internet_search_results.json', 'w') as f:
                json.dump(parent.internet_search_results, f, indent=2)
        except Exception as e:
            messagebox.showerror("Error", f"Error saving internet search results: {e}")

    def save_texts(self, texts, filename):
        try:
            with open(filename, 'w') as f:
                json.dump(texts, f, indent=2)
        except Exception as e:
            messagebox.showerror("Error", f"Error saving texts: {e}")

    def format_scripts(self, scripts):
        return "\n\n".join([f"Script {i+1} ({name}):\n{content}\n!!!this is the next document!!!" for i, (name, content) in enumerate(scripts)])

    def format_instructions(self, instructions):
        return "\n\n".join([f"Instruction {i+1} ({name}):\n{content}\n!!!this is the next document!!!" for i, (name, content) in enumerate(instructions)])

    def format_internet_sources(self, internet_sources):
        formatted_sources = []
        for i, source in enumerate(internet_sources):
            formatted_sources.append(
                f"Internet Source {i+1} (URL: {source['url']}, Author: {source['author']}, Date: {source['date']}):\n{source['content']}\n!!!this is the next document!!!"
            )
        return "\n\n".join(formatted_sources)

    def format_internet_search_results(self, internet_search_results):
        formatted_results = []
        for i, result in enumerate(internet_search_results):
            formatted_results.append(
                f"Internet Search Result {i+1} (Title: {result.get('title', 'Unknown')}, URL: {result.get('url', 'unknown')}, Author: {result.get('author', 'Unknown')}, Date Retrieved: {result.get('date_retrieved', 'N/A')}):\n{result.get('content', 'No content available')}\n!!!this is the next document!!!"
            )
        return "\n\n".join(formatted_results)

    def load_all_settings(self):
        try:
            with open('claude_app_settings.json', 'r') as f:
                settings = json.load(f)
        except FileNotFoundError:
            settings = {}
        
        try:
            with open('script_texts.json', 'r') as f:
                settings['scripts'] = json.load(f)
        except FileNotFoundError:
            settings['scripts'] = []
        
        try:
            with open('instruction_texts.json', 'r') as f:
                settings['instructions'] = json.load(f)
        except FileNotFoundError:
            settings['instructions'] = []
        
        try:
            with open('internet_sources.json', 'r') as f:
                settings['internet_sources'] = json.load(f)
        except FileNotFoundError:
            settings['internet_sources'] = []
        
        try:
            with open('internet_search_results.json', 'r') as f:
                settings['internet_search_results'] = json.load(f)
        except FileNotFoundError:
            settings['internet_search_results'] = []
        
        return settings

    def save_all_settings(self, parent):
        settings = {
            'api_key': parent.api_key,
            'perplexity_api_key': parent.perplexity_api_key,
            'first_name': parent.first_name,
            'last_name': parent.last_name,
            'date': parent.date,
            'font_name': parent.font_name,
            'font_size_normal': parent.font_size_normal,
            'font_size_heading1': parent.font_size_heading1,
            'font_size_heading2': parent.font_size_heading2,
            'font_size_heading3': parent.font_size_heading3,
            'line_spacing': parent.line_spacing,
            'margin_top': parent.margin_top,
            'margin_bottom': parent.margin_bottom,
            'margin_left': parent.margin_left,
            'margin_right': parent.margin_right,
            'system_prompt': parent.system_prompt_text.get(1.0, tk.END).strip(),
            'custom_prompts': parent.custom_prompts,
            'scripts': parent.scripts,
            'instructions': parent.instructions,
            'internet_sources': parent.internet_sources,
            'internet_search_results': parent.internet_search_results
        }
        try:
            with open('claude_app_settings.json', 'w') as f:
                json.dump(settings, f)
        except Exception as e:
            messagebox.showerror("Error", f"Error saving settings: {e}")

    def save_custom_prompt(self, parent):
        prompt_name = tk.simpledialog.askstring("Save Prompt", "Enter a name for this prompt:")
        if prompt_name:
            current_prompt = parent.system_prompt_text.get(1.0, tk.END).strip()
            parent.custom_prompts[prompt_name] = current_prompt
            parent.save_all_settings()
            messagebox.showinfo("Success", f"Prompt '{prompt_name}' saved successfully.")



class APIHandler:
        def send_request(self, parent):
            if not parent.instructions:
                messagebox.showerror("Error", "Please upload instruction files first.")
                return
            if not parent.scripts:
                messagebox.showerror("Error", "Please upload script files first.")
                return
            if not parent.first_name or not parent.last_name:
                messagebox.showerror("Error", "Please enter your first and last name in the settings.")
                return
            if not parent.api_key:
                messagebox.showerror("Error", "Please enter your API key in the settings.")
                return

            parent.update_system_prompt()

            messages = [
                {"role": "user", "content": parent.system_prompt}
            ]

            api_url = "https://api.anthropic.com/v1/messages"
            headers = {
                "x-api-key": parent.api_key,
                "anthropic-version": "2023-06-01",
                "content-type": "application/json",
                "anthropic-beta": "max-tokens-3-5-sonnet-2024-07-15"
            }
            data = {
                "model": "claude-3-5-sonnet-20240620",
                "max_tokens": 8192,
                "messages": messages
            }

            try:
                parent.output_text.delete(1.0, tk.END)
                parent.output_text.insert(tk.END, "Generating response, please wait...")
                parent.update_idletasks()

                response = requests.post(api_url, headers=headers, json=data)

                if response.status_code == 200:
                    result = response.json()
                    content = result['content'][0]['text']
                    response_text = content.strip()
                    parent.output_text.delete(1.0, tk.END)
                    parent.output_text.insert(tk.END, response_text)
                    messagebox.showinfo("Success", "Paper generated.")
                else:
                    error_message = response.text
                    parent.output_text.delete(1.0, tk.END)
                    messagebox.showerror("Error", f"API Error {response.status_code}: {error_message}")
            except Exception as e:
                parent.output_text.delete(1.0, tk.END)
                messagebox.showerror("Error", f"Error making API request: {e}")

class DocumentHandler:
    def save_output(self, parent):
        output = parent.output_text.get(1.0, tk.END).strip()
        if not output:
            messagebox.showerror("Error", "No output to save.")
            return

        save_path = filedialog.asksaveasfilename(title="Save Output as Word File", defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if save_path:
            try:
                document = Document()
                self.set_document_properties(document, parent)
                self.process_content(document, output, parent)
                self.add_page_numbers(document.sections[0])
                document.save(save_path)
                messagebox.showinfo("Success", f"Output saved to {save_path}")
            except Exception as e:
                messagebox.showerror("Error", f"Error saving Word file: {e}")

    def set_document_properties(self, document, parent):
        section = document.sections[0]
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        section.left_margin = Cm(parent.margin_left)
        section.right_margin = Cm(parent.margin_right)
        section.top_margin = Cm(parent.margin_top)
        section.bottom_margin = Cm(parent.margin_bottom)

        styles = document.styles
        self.create_custom_style(styles, 'TitleStyle', parent.font_size_heading1, True, parent)
        self.create_custom_style(styles, 'Heading1Custom', parent.font_size_heading1, True, parent, base_style='Heading 1')
        self.create_custom_style(styles, 'Heading2Custom', parent.font_size_heading2, True, parent, base_style='Heading 2')
        self.create_custom_style(styles, 'Heading3Custom', parent.font_size_heading3, True, parent, base_style='Heading 3')

        style_normal = styles['Normal']
        style_normal.font.size = Pt(parent.font_size_normal)
        style_normal.font.name = parent.font_name

        line_spacing = 1.0 if parent.line_spacing == "Single" else 1.5 if parent.line_spacing == "1.5 lines" else 2.0
        for style in [style_normal, styles['TitleStyle'], styles['Heading1Custom'], styles['Heading2Custom'], styles['Heading3Custom']]:
            style.paragraph_format.line_spacing = Pt(line_spacing * 12)

    def create_custom_style(self, styles, name, font_size, bold, parent, base_style=None):
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        if base_style:
            style.base_style = styles[base_style]
        style.font.size = Pt(font_size)
        style.font.bold = bold
        style.font.name = parent.font_name

    def process_content(self, document, content, parent):
        paragraphs = content.strip().split('\n')
        i = 0
        while i < len(paragraphs):
            para = paragraphs[i].strip()
            if not para:
                i += 1
                continue

            if para == '####TITLE PAGE####':
                i = self.process_title_page(document, paragraphs, i)
            elif para.startswith('# '):
                document.add_paragraph(para[2:].strip(), style='Heading1Custom')
            elif para.startswith('## '):
                document.add_paragraph(para[3:].strip(), style='Heading2Custom')
            elif para.startswith('### '):
                document.add_paragraph(para[4:].strip(), style='Heading3Custom')
            elif re.match(r'^\d+\.', para) or para.startswith('- '):
                p = document.add_paragraph(style='List Bullet')
                self.add_runs(p, para.lstrip('0123456789.- '))
            elif para.startswith('|') and para.endswith('|'):
                i = self.process_table(document, paragraphs, i, parent)
            else:
                p = document.add_paragraph(style='Normal')
                self.add_runs(p, para)
            i += 1

    def process_title_page(self, document, paragraphs, i):
        title_page_content = []
        i += 1
        while i < len(paragraphs) and paragraphs[i].strip() != '####END TITLE PAGE####':
            title_page_content.append(paragraphs[i].strip())
            i += 1
        for line in title_page_content:
            if line.startswith('#'):
                p = document.add_paragraph(line.lstrip('#').strip(), style='TitleStyle')
            else:
                p = document.add_paragraph(line, style='Normal')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_page_break()
        return i + 1

    def process_table(self, document, paragraphs, i, parent):
        table_lines = [paragraphs[i]]
        i += 1
        while i < len(paragraphs) and paragraphs[i].strip().startswith('|') and paragraphs[i].strip().endswith('|'):
            table_lines.append(paragraphs[i].strip())
            i += 1
        table = self.parse_markdown_table(table_lines)
        if table:
            word_table = document.add_table(rows=len(table), cols=len(table[0]))
            word_table.style = 'Table Grid'
            for row_idx, row in enumerate(table):
                for col_idx, cell in enumerate(row):
                    word_table.cell(row_idx, col_idx).text = cell
            for row in word_table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.style = document.styles['Normal']
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return i - 1

    def parse_markdown_table(self, table_lines):
        try:
            table = [list(filter(None, line.strip('|').split('|'))) for line in table_lines]
            if len(table) > 1 and all(cell.strip().startswith('-') for cell in table[1]):
                table.pop(1)
            return table
        except Exception:
            return None

    def add_runs(self, paragraph, text):
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                paragraph.add_run(part)

    def add_page_numbers(self, section):
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        fldSimple = OxmlElement('w:fldSimple')
        fldSimple.set(qn('w:instr'), 'PAGE')
        run._r.append(fldSimple)
